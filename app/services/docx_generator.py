import io
import re
import matplotlib
# Set non-interactive backend before importing pyplot
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import List, Dict, Any

class DocxGenerator:
    def __init__(self):
        self.doc = Document()

    def generate_report(self, title: str, summary: str, run_data: List[Dict[str, Any]], metrics_data: List[Dict[str, Any]]) -> io.BytesIO:
        """
        Generates a Word document report with charts.
        
        Args:
            title: Report title
            summary: Text summary
            run_data: List of dicts with 'version', 'score', 'created_at', 'gap_analysis'
            metrics_data: List of dicts with 'metric_name', 'scores': [{'version': v, 'score': s}]
        """
        self._add_title(title)
        self._add_summary(summary)
        
        # Add Charts
        self.doc.add_heading('Performance Visualizations', level=1)
        
        # 1. Aggregated Score Over Time
        score_chart = self._create_score_chart(run_data)
        if score_chart:
            self.doc.add_paragraph("Aggregated Score Progression")
            self.doc.add_picture(score_chart, width=Inches(6))
            score_chart.close()

        # 2. Metric Breakdown (Bar chart of latest version or line chart of all? Let's do line chart for top metrics)
        metric_chart = self._create_metric_chart(run_data, metrics_data)
        if metric_chart:
            self.doc.add_paragraph("Metric-Specific Trends")
            self.doc.add_picture(metric_chart, width=Inches(6))
            metric_chart.close()
            
        self.doc.add_page_break()
        
        # Add Detailed Analysis
        self.doc.add_heading('Detailed Progress Analysis', level=1)
        
        # Reverse order to show latest first? Or chronological? 
        # Requirement: "analyze the gap reports... and come up with detailed text".
        # The 'summary' arg contains the AI generated meta-analysis.
        # But we also want to show the per-version details as "Evidence".
        
        for run in sorted(run_data, key=lambda x: x['version']):
            self.doc.add_heading(f"Version {run['version']} (Score: {run['score']})", level=2)
            self.doc.add_paragraph(f"Date: {run['created_at'].strftime('%Y-%m-%d %H:%M')}")
            
            if run.get('gap_analysis'):
                self.doc.add_heading('Gap Analysis', level=3)
                self.doc.add_paragraph(run['gap_analysis'])
            else:
                self.doc.add_paragraph("No gap analysis available.", style='Italic')
                
        # Save to buffer
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _add_title(self, text: str):
        heading = self.doc.add_heading(text, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_summary(self, text: str):
        self.doc.add_heading('Executive Summary', level=1)
        # Parse paragraphs (double newline)
        paragraphs = text.split('\n\n')
        for p_text in paragraphs:
            self._add_markdown_paragraph(p_text.strip())

    def _add_markdown_paragraph(self, text: str):
        if not text:
            return
        
        p = self.doc.add_paragraph()
        # Regex to split by bold (**...**) and italic (*...*)
        # Keep delimiters to know what matches
        # Note: This is a simple parser. Nested styles not supported.
        # Order: Bold then Italic.
        
        # Tokenize by bold **
        tokens = re.split(r'(\*\*[^*]+\*\*)', text)
        for token in tokens:
            if token.startswith('**') and token.endswith('**') and len(token) > 4:
                run = p.add_run(token[2:-2])
                run.bold = True
            else:
                # Tokenize by italic * (if not inside bold, but splitting handled that)
                # It might catch single * inside text, so be careful. 
                # Simplistic: look for *...*
                sub_tokens = re.split(r'(\*[^*]+\*)', token)
                for sub_token in sub_tokens:
                    if sub_token.startswith('*') and sub_token.endswith('*') and len(sub_token) > 2:
                        run = p.add_run(sub_token[1:-1])
                        run.italic = True
                    else:
                        p.add_run(sub_token)

    def _create_score_chart(self, run_data: List[Dict]):
        if not run_data:
            return None
            
        versions = [str(r['version']) for r in run_data]
        scores = [r['score'] or 0 for r in run_data]
        
        plt.figure(figsize=(10, 5))
        plt.plot(versions, scores, marker='o', linestyle='-', color='#003366', linewidth=2)
        plt.title('Aggregated Score by Version')
        plt.xlabel('Version')
        plt.ylabel('Score')
        plt.grid(True, linestyle='--', alpha=0.7)
        plt.ylim(0, 100) # Assuming 0-100 scale
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf

    def _create_metric_chart(self, run_data: List[Dict], metrics_data: List[Dict]):
        if not metrics_data:
            return None
            
        plt.figure(figsize=(10, 5))
        
        # Extract versions for X axis
        all_versions = sorted([r['version'] for r in run_data])
        str_versions = [str(v) for v in all_versions]
        
        # Plot top 5 metrics to avoid clutter
        for metric in metrics_data[:5]:
            # Align scores to versions
            m_scores = []
            score_map = {s['version']: s['score'] for s in metric['scores']}
            for v in all_versions:
                m_scores.append(score_map.get(v, None)) # Handle missing data with gaps? Matplotlib handles None? 
                
            # Filter None for plotting if needed, or rely on plot gap handling
            # Matplotlib plot with None breaks the line, which is correct.
            plt.plot(str_versions, m_scores, marker='x', linestyle='--', label=metric['metric_name'])
            
        plt.title('Individual Metric Performance')
        plt.xlabel('Version')
        plt.ylabel('Score')
        plt.legend()
        plt.grid(True, alpha=0.3)
        plt.ylim(0, 100)
        
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf

def generate_word_report(title: str, summary: str, run_data: List[Dict], metrics_data: List[Dict]) -> io.BytesIO:
    generator = DocxGenerator()
    return generator.generate_report(title, summary, run_data, metrics_data)
